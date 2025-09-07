import uuid

from flask import Flask, render_template, request, jsonify
from flask_socketio import SocketIO, emit
import openai
import tempfile
import os
import json
import base64
import logging
from supabase import create_client, Client
import datetime
from PyPDF2 import PdfReader
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

# Настройка логирования
logging.basicConfig(level=logging.INFO)

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-here'
socketio = SocketIO(app, cors_allowed_origins="*")

# OpenAI ключ
openai.api_key = ""

# Supabase конфигурация
SUPABASE_URL = ""
SUPABASE_KEY = ""
supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

# Конфигурация
VOICE = "nova"
MODEL_TTS = "tts-1"
UPLOAD_FOLDER = 'uploads'
AUDIO_FOLDER = 'audio'
VACANCIES_FOLDER = 'vacancies'
INTERVIEW_DURATION_LIMIT_MINUTES = 30
INTERVIEW_MESSAGE_LIMIT = 20

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(AUDIO_FOLDER, exist_ok=True)
os.makedirs(VACANCIES_FOLDER, exist_ok=True)

sessions = {}


class HRBotSession:
    def __init__(self, session_id, vacancy_id, candidate_id):
        self.session_id = session_id
        self.vacancy_id = vacancy_id
        self.candidate_id = candidate_id
        self.resume_text = ""
        self.requirements_text = ""
        self.resume_summary = ""
        self.requirements_summary = ""
        self.conversation_history = []
        self.start_time = None
        self.end_time = None
        self.rudeness_detected = False
        self.skill_mismatch_count = 0


def get_session(session_id):
    if session_id not in sessions:
        data = supabase.table('candidates').select('*').eq('id', session_id).execute().data
        if not data:
            raise ValueError("Кандидат не найден")
        candidate = data[0]
        sessions[session_id] = HRBotSession(session_id, candidate['vacancy_id'], session_id)
    return sessions[session_id]


def read_file_text(file_path):
    try:
        if file_path.lower().endswith('.pdf'):
            reader = PdfReader(file_path)
            text = "".join([page.extract_text() or "" for page in reader.pages]).strip()
            if not text:
                raise ValueError("Не удалось извлечь текст из PDF.")
            return text
        elif file_path.lower().endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read().strip()
            if not text:
                raise ValueError("TXT файл пустой.")
            return text
        elif file_path.lower().endswith('.docx'):
            try:
                doc = Document(file_path)
                text = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        text.append(para.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text.append(cell.text.strip())
                for section in doc.sections:
                    for header in section.header.paragraphs:
                        if header.text.strip():
                            text.append(header.text.strip())
                    for footer in section.footer.paragraphs:
                        if footer.text.strip():
                            text.append(footer.text.strip())
                text = "\n".join(text).strip()
                if not text:
                    raise ValueError("DOCX файл не содержит извлекаемого текста.")
                return text
            except PackageNotFoundError:
                raise ValueError("DOCX файл повреждён или не является действительным DOCX.")
        else:
            raise ValueError("Поддерживаемые форматы: PDF, TXT или DOCX")
    except Exception as e:
        raise ValueError(f"Ошибка обработки файла: {str(e)}")


def analyze_sentiment_and_skills(user_input, resume_summary, requirements_summary):
    prompt = (
        f"Анализируй сообщение кандидата на предмет грубости и соответствия навыкам:\n"
        f"Сообщение: {user_input}\n"
        f"Резюме: {resume_summary}\n"
        f"Требования вакансии: {requirements_summary}\n"
        f"1. Определи, содержит ли сообщение грубость или неуважительное поведение (например, оскорбления, сарказм). "
        f"Ответь 'грубость обнаружена' или 'грубость не обнаружена'.\n"
        f"2. Проверь, подтверждает ли сообщение навыки, указанные в требованиях вакансии. "
        f"Ответь 'навыки подтверждены', 'навыки не подтверждены' или 'навыки частично подтверждены'.\n"
        f"Формат ответа:\n- Сентимент: [ответ]\n- Навыки: [ответ]"
    )
    try:
        response = openai.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5
        )
        result = response.choices[0].message.content.strip().split('\n')
        sentiment = result[0].split(': ')[1]
        skills = result[1].split(': ')[1]
        return sentiment, skills
    except Exception as e:
        return "ошибка анализа", "ошибка анализа"


def generate_summary(text, summary_type="resume"):
    prompt = (
        f"Анализируй это резюме кандидата и извлеки ключевые моменты:\n- Имя, - Навыки, - Опыт, - Образование, - Проекты.\nСделай краткое саммари на русском с маркерами.\nТекст: {text}"
        if summary_type == "resume" else
        f"Анализируй требования к вакансии: - Навыки, - Опыт, - Образование, - Другие.\nСделай краткое саммари на русском с маркерами.\nТекст: {text}"
    )
    try:
        response = openai.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": prompt}],
                                                  temperature=0.5)
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Ошибка генерации саммари: {str(e)}"


def generate_response(user_input, resume_summary, requirements_summary, conversation_history, session):
    sentiment, skills = analyze_sentiment_and_skills(user_input, resume_summary, requirements_summary)

    if sentiment == "грубость обнаружена":
        session.rudeness_detected = True
        return (
            "Ваше поведение недопустимо. Интервью завершено из-за проявления неуважения.",
            True  # Indicates interview should end
        )

    if skills == "навыки не подтверждены":
        session.skill_mismatch_count += 1
        if session.skill_mismatch_count >= 3:  # Allow 3 failed skill confirmations
            return (
                "Ваши ответы не подтверждают необходимые навыки для этой вакансии. Интервью завершено.",
                True
            )

    context = "\n".join([f"Пользователь: {entry['user']}\nБот: {entry['bot']}" for entry in
                         conversation_history[-5:]]) if conversation_history else ""
    prompt = (
        f"Ты HR-бот Анна. Будь жёсткой. Контекст: Резюме: {resume_summary}\nТребования: {requirements_summary}\n{context}\n"
        f"Пользователь: {user_input}\nОтветь естественно на русском, задавай вопросы по одному, чтобы проверить навыки."
    )
    try:
        response = openai.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": prompt}],
                                                  temperature=0.7)
        return response.choices[0].message.content.strip(), False
    except Exception as e:
        return f"Ошибка генерации ответа: {str(e)}", False


def calculate_match_percentage(resume_summary, requirements_summary, conversation_history):
    prompt = (
        f"Оцени соответствие кандидата требованиям вакансии в процентах (0-100):\n"
        f"Резюме: {resume_summary}\nТребования: {requirements_summary}\n"
        f"История диалога: {'\n'.join([f'Пользователь: {entry['user']}\nБот: {entry['bot']}' for entry in conversation_history])}\n"
        f"Верни только число от 0 до 100."
    )
    try:
        response = openai.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": prompt}],
                                                  temperature=0.5)
        percentage = int(response.choices[0].message.content.strip())
        return max(0, min(100, percentage))  # Ensure within 0-100
    except Exception:
        return 50  # Fallback to neutral score


def generate_interview_summary(resume_summary, requirements_summary, conversation_history, rudeness_detected,
                               skill_mismatch_count):
    history_text = "\n".join([f"Пользователь: {entry['user']}\nБот: {entry['bot']}" for entry in conversation_history])
    prompt = (
        f"Анализируй: Резюме: {resume_summary}\nТребования: {requirements_summary}\nИстория: {history_text}\n"
        f"Грубость обнаружена: {'да' if rudeness_detected else 'нет'}\n"
        f"Количество неподтверждённых навыков: {skill_mismatch_count}\n"
        f"Оцени соответствие: навыки, сильные/слабые стороны, рекомендация. Укажи причину завершения, если интервью было прервано."
    )
    try:
        response = openai.chat.completions.create(model="gpt-4o-mini", messages=[{"role": "user", "content": prompt}],
                                                  temperature=0.5)
        summary = response.choices[0].message.content.strip()
        if rudeness_detected:
            summary += "\nИнтервью завершено досрочно из-за неуважительного поведения кандидата."
        elif skill_mismatch_count >= 3:
            summary += "\nИнтервью завершено досрочно из-за неподтверждения необходимых навыков."
        return summary
    except Exception as e:
        return f"Ошибка генерации саммари интервью: {str(e)}"


def transcribe_audio(audio_data):
    with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as f:
        f.write(audio_data)
        temp_filename = f.name
    try:
        with open(temp_filename, "rb") as f:
            transcript = openai.audio.transcriptions.create(model="whisper-1", file=f)
        os.unlink(temp_filename)
        return transcript.text.strip()
    except Exception as e:
        os.unlink(temp_filename)
        raise ValueError(f"Ошибка транскрипции: {str(e)}")


def text_to_speech(text):
    try:
        response = openai.audio.speech.create(model=MODEL_TTS, voice=VOICE, input=text)
        with tempfile.NamedTemporaryFile(suffix=".mp3", delete=False) as f:
            response.stream_to_file(f.name)
            return f.name
    except Exception as e:
        raise ValueError(f"Ошибка генерации аудио: {str(e)}")


# Роуты для страниц
@app.route('/')
def index():
    return "Добро пожаловать в HR-Бот! Выберите: <a href='/employer'>Работодатель</a> или <a href='/candidate'>Соискатель</a>"


@app.route('/employer')
def employer():
    return render_template('employer_tg.html')


@app.route('/candidate')
def candidate():
    return render_template('candidate_tg.html')


# Роуты для данных
@app.route('/get_vacancies', methods=['GET'])
def get_vacancies():
    try:
        data = supabase.table('vacancies').select('*').execute().data
        return jsonify(data or [])
    except Exception as e:
        return jsonify({'error': f"Ошибка загрузки вакансий: {str(e)}"}), 500


@app.route('/get_candidates/<vacancy_id>', methods=['GET'])
def get_candidates(vacancy_id):
    try:
        vacancy_id = int(vacancy_id)
        data = supabase.table('candidates').select('*').eq('vacancy_id', vacancy_id).execute().data
        return jsonify(data or [])
    except ValueError:
        return jsonify({'error': 'Неверный ID вакансии'}), 400
    except Exception as e:
        return jsonify({'error': f"Ошибка загрузки кандидатов: {str(e)}"}), 500


@app.route('/get_schedules/<candidate_id>', methods=['GET'])
def get_schedules(candidate_id):
    try:
        candidate_id = int(candidate_id)
        data = supabase.table('schedules').select('*').eq('candidate_id', candidate_id).execute().data
        return jsonify(data or [])
    except ValueError:
        return jsonify({'error': 'Неверный ID кандидата'}), 400
    except Exception as e:
        return jsonify({'error': f"Ошибка загрузки расписания: {str(e)}"}), 500


@app.route('/get_interview/<candidate_id>', methods=['GET'])
def get_interview(candidate_id):
    try:
        candidate_id = int(candidate_id)
        data = supabase.table('interviews').select('*').eq('candidate_id', candidate_id).execute().data
        return jsonify(data or [])
    except ValueError:
        return jsonify({'error': 'Неверный ID кандидата'}), 400
    except Exception as e:
        return jsonify({'error': f"Ошибка загрузки интервью: {str(e)}"}), 500


@app.route('/create_vacancy', methods=['POST'])
def create_vacancy():
    try:
        data = request.get_json()
        title = data.get('title', 'Untitled Vacancy')
        insert = supabase.table('vacancies').insert({'title': title}).execute().data[0]
        vacancy_id = insert['id']
        folder_path = os.path.join(VACANCIES_FOLDER, str(vacancy_id))
        os.makedirs(folder_path, exist_ok=True)
        supabase.table('vacancies').update({'folder_path': folder_path}).eq('id', vacancy_id).execute()
        return jsonify({'success': True, 'vacancy_id': vacancy_id, 'folder_path': folder_path})
    except Exception as e:
        return jsonify({'error': f"Ошибка создания вакансии: {str(e)}"}), 500


@app.route('/upload', methods=['POST'])
def upload():
    try:
        vacancy_id = request.form.get('vacancy_id')
        file_type = request.form.get('type')
        file = request.files.get('file')
        if not vacancy_id or vacancy_id == 'null' or not file_type or not file:
            return jsonify({'error': 'Неверные данные или ID вакансии'}), 400
        vacancy_id = int(vacancy_id)
        vacancy_data = supabase.table('vacancies').select('*').eq('id', vacancy_id).execute().data
        if not vacancy_data:
            return jsonify({'error': 'Вакансия не найдена'}), 400
        vacancy = vacancy_data[0]
        folder_path = vacancy['folder_path']
        os.makedirs(folder_path, exist_ok=True)
        safe_filename = "".join(c for c in file.filename if c.isalnum() or c in ('.', '_')).rstrip()
        filename = f"{file_type}_{safe_filename}"
        filepath = os.path.join(folder_path, filename)
        file.save(filepath)
        text = read_file_text(filepath)
        summary = generate_summary(text, file_type)
        if file_type == 'requirements':
            supabase.table('vacancies').update({'requirements_text': text, 'requirements_summary': summary}).eq('id',
                                                                                                                vacancy_id).execute()
        elif file_type == 'resume':
            insert = supabase.table('candidates').insert(
                {'name': safe_filename.split('.')[0], 'resume_text': text, 'resume_summary': summary,
                 'vacancy_id': vacancy_id}).execute().data[0]
            candidate_id = insert['id']
            schedule_time = datetime.datetime.now() + datetime.timedelta(days=1)
            supabase.table('schedules').insert(
                {'candidate_id': candidate_id, 'interview_time': schedule_time.isoformat()}).execute()
            return jsonify({'success': True, 'candidate_id': candidate_id, 'summary': summary})
        return jsonify({'success': True, 'summary': summary})
    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        return jsonify({'error': f"Ошибка загрузки файла: {str(e)}"}), 500


@app.route('/chat', methods=['POST'])
def chat():
    try:
        data = request.get_json()
        session_id = data.get('session_id')
        user_message = data.get('message', '')
        session = get_session(session_id)
        if session.start_time and (
                datetime.datetime.now() - session.start_time).total_seconds() / 60 > INTERVIEW_DURATION_LIMIT_MINUTES or len(
                session.conversation_history) >= INTERVIEW_MESSAGE_LIMIT:
            return jsonify({'error': 'Интервью завершено'}), 400
        bot_response, should_end = generate_response(user_message, session.resume_summary, session.requirements_summary,
                                                     session.conversation_history, session)
        session.conversation_history.append({'user': user_message, 'bot': bot_response})
        if should_end:
            match_percentage = calculate_match_percentage(session.resume_summary, session.requirements_summary,
                                                          session.conversation_history)
            summary = generate_interview_summary(session.resume_summary, session.requirements_summary,
                                                 session.conversation_history, session.rudeness_detected,
                                                 session.skill_mismatch_count)
            supabase.table('interviews').insert({
                'candidate_id': session_id,
                'start_time': session.start_time.isoformat(),
                'end_time': datetime.datetime.now().isoformat(),
                'conversation_history': json.dumps(session.conversation_history),
                'summary': summary,
                'match_percentage': match_percentage
            }).execute()
            if session_id in sessions:
                del sessions[session_id]
            return jsonify({'response': bot_response, 'end_interview': True, 'summary': summary,
                            'match_percentage': match_percentage})
        audio_file = text_to_speech(bot_response)
        with open(audio_file, 'rb') as f:
            audio_data = base64.b64encode(f.read()).decode('utf-8')
        os.unlink(audio_file)
        return jsonify({'response': bot_response, 'audio': audio_data})
    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        return jsonify({'error': f"Ошибка обработки чата: {str(e)}"}), 500


@app.route('/transcribe', methods=['POST'])
def transcribe():
    try:
        audio_file = request.files.get('audio')
        if not audio_file:
            return jsonify({'error': 'Аудио не найдено'}), 400
        audio_data = audio_file.read()
        text = transcribe_audio(audio_data)
        return jsonify({'transcription': text})
    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        return jsonify({'error': f"Ошибка транскрипции: {str(e)}"}), 500


# Socket.IO
@socketio.on('connect')
def handle_connect():
    emit('connected', {'message': 'Подключено'})


@socketio.on('start_interview')
def handle_start_interview(data):
    session_id = data.get('session_id')
    try:
        session = get_session(session_id)
        candidate_data = supabase.table('candidates').select('*').eq('id', session_id).execute().data[0]
        vacancy_data = supabase.table('vacancies').select('*').eq('id', candidate_data['vacancy_id']).execute().data[0]
        session.resume_summary = candidate_data['resume_summary'] or ""
        session.requirements_summary = vacancy_data['requirements_summary'] or ""
        session.start_time = datetime.datetime.now()
        initial_message = "Привет! Меня зовут Анна. Расскажите о вашем опыте."
        audio_file = text_to_speech(initial_message)
        with open(audio_file, 'rb') as f:
            audio_data = base64.b64encode(f.read()).decode('utf-8')
        os.unlink(audio_file)
        session.conversation_history.append({'user': '', 'bot': initial_message})
        emit('interview_started', {'message': initial_message, 'audio': audio_data})
    except Exception as e:
        emit('error', {'message': f"Ошибка старта интервью: {str(e)}"})


@socketio.on('end_interview')
def handle_end_interview(data):
    session_id = data.get('session_id')
    try:
        session = get_session(session_id)
        session.end_time = datetime.datetime.now()
        match_percentage = calculate_match_percentage(session.resume_summary, session.requirements_summary,
                                                      session.conversation_history)
        summary = generate_interview_summary(session.resume_summary, session.requirements_summary,
                                             session.conversation_history, session.rudeness_detected,
                                             session.skill_mismatch_count)
        supabase.table('interviews').insert({
            'candidate_id': session_id,
            'start_time': session.start_time.isoformat(),
            'end_time': session.end_time.isoformat(),
            'conversation_history': json.dumps(session.conversation_history),
            'summary': summary,
            'match_percentage': match_percentage
        }).execute()
        if session_id in sessions:
            del sessions[session_id]
        emit('interview_ended', {'summary': summary, 'match_percentage': match_percentage})
    except Exception as e:
        emit('error', {'message': f"Ошибка завершения интервью: {str(e)}"})

@app.route('/regenerate_auth_code/<candidate_id>', methods=['POST'])
def regenerate_auth_code(candidate_id):
    try:
        logging.debug(f"Attempting to regenerate auth code for candidate_id: {candidate_id}")
        candidate_id = int(candidate_id)
        candidate_data = supabase.table('candidates').select('id').eq('id', candidate_id).execute().data
        if not candidate_data:
            logging.warning(f"Candidate with id {candidate_id} not found")
            return jsonify({'error': 'Кандидат не найден'}), 404
        new_auth_code = str(uuid.uuid4())
        supabase.table('candidates').update({
            'auth_code': new_auth_code,
            'is_auth_code_used': False
        }).eq('id', candidate_id).execute()
        logging.info(f"New auth code generated for candidate_id {candidate_id}: {new_auth_code}")
        return jsonify({'success': True, 'new_auth_code': new_auth_code}), 200
    except ValueError:
        logging.error(f"Invalid candidate_id: {candidate_id}")
        return jsonify({'error': 'Неверный ID кандидата: должен быть числом'}), 400
    except Exception as e:
        logging.error(f"Error regenerating auth code: {str(e)}")
        return jsonify({'error': f"Внутренняя ошибка сервера: {str(e)}"}), 500

if __name__ == '__main__':
    socketio.run(app, debug=True, host='0.0.0.0', port=5000)